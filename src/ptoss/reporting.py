from __future__ import annotations

import os
from hashlib import sha256
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
except ImportError:  # pragma: no cover - fallback for minimal environments
    Document = None

from .models import AssessmentCase, ReportArtifact, new_id, utc_now
from .rules import DRSResult


REPORT_ROOT = Path(os.getenv("PTOSS_REPORT_DIR", "data/reports"))


def _report_lines(case: AssessmentCase, drs_result: DRSResult, engine_version: str) -> list[str]:
    generated_at = utc_now()
    lines: list[str] = []
    lines.append("# PT-OSS HealthCheck Studio Provisional Report")
    lines.append("")
    lines.append(f"- Case ID: {case.id}")
    lines.append(f"- Organization ID: {case.organization.id}")
    lines.append(f"- Organization: {case.organization.name}")
    lines.append(f"- Assessment Mode: {case.scope.assessment_mode.value}")
    lines.append(f"- Generated At: {generated_at.isoformat()}")
    lines.append(f"- Rule Package Version: {case.rule_package_version}")
    lines.append(f"- Evidence Index Version: {case.evidence_index_version}")
    lines.append(f"- Engine Version: {engine_version}")
    lines.append("- Reviewer Status: REVIEWED BY BOARD, PROVISIONAL OUTPUT")
    lines.append("- Official Label: PROVISIONAL")
    lines.append("- Verification ID: pending")
    lines.append(f"- DRS: {drs_result.value}")
    lines.append(f"- Band: {drs_result.band.value}")
    lines.append("")
    lines.append("## Evidence Gaps")
    if drs_result.critical_alerts:
        for alert in drs_result.critical_alerts:
            lines.append(f"- {alert.rule_code}: {alert.message}")
    else:
        lines.append("- None")
    lines.append("")
    lines.append("## 24/72-hour Escalation Plan")
    lines.append("- 24h: confirm critical evidence gaps and assign owners.")
    lines.append("- 72h: close missing items or record explicit justification.")
    lines.append("")
    lines.append("## 30/60/90-day Treatment Plan")
    lines.append("- 30d: stabilize missing core documents and review owners.")
    lines.append("- 60d: verify repeatability and quality of evidence updates.")
    lines.append("- 90d: re-assess readiness and refresh the case baseline.")
    lines.append("")
    lines.append("## Status")
    lines.append("PROVISIONAL - Not Reviewed, Not Certified")
    return lines


def _write_docx_report(path: Path, lines: Iterable[str]) -> None:
    if Document is None:  # pragma: no cover - fallback for minimal environments
        path.write_text("\n".join(lines), encoding="utf-8")
        return

    document = Document()
    document.add_heading("PT-OSS HealthCheck Studio Provisional Report", level=0)
    current_heading = None
    for line in lines:
        if line.startswith("## "):
            current_heading = line[3:]
            document.add_heading(current_heading, level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif not line.strip():
            document.add_paragraph("")
        elif line.startswith("# "):
            continue
        else:
            document.add_paragraph(line)
    document.save(path)


def render_provisional_report(case: AssessmentCase, drs_result: DRSResult, engine_version: str) -> ReportArtifact:
    lines = _report_lines(case, drs_result, engine_version)
    content = "\n".join(lines)
    content_hash = sha256(content.encode("utf-8")).hexdigest()
    report_id = new_id("report")
    report_dir = REPORT_ROOT / case.id
    report_dir.mkdir(parents=True, exist_ok=True)
    docx_path = report_dir / f"{report_id}.docx"
    md_path = report_dir / f"{report_id}.md"
    md_path.write_text(content, encoding="utf-8")
    _write_docx_report(docx_path, lines)
    return ReportArtifact(
        id=report_id,
        case_id=case.id,
        tenant_id=case.tenant_id,
        report_type="HEALTH_CHECK_REPORT",
        file_format="DOCX",
        artifact_uri=str(docx_path.resolve()),
        official_status="PROVISIONAL",
        engine_version=engine_version,
        rule_package_version=case.rule_package_version,
        evidence_index_version=case.evidence_index_version,
        verification_id=None,
        content_hash=content_hash,
        content=content,
    )
